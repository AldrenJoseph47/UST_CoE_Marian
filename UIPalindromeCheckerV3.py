import tkinter as tk
from tkinter import ttk
import os.path
import csv
from tkinter import messagebox
from docx import Document
import shutil
from tkinter import filedialog #importing filer in the csv of doc format

window = tk.Tk()
window.geometry("640x400")
window.title("PALINDROME CHECKER")
window.resizable(0, 0)
window['bg'] = "#CAA8F5"

#naming the csv and doc files variables
txt_file = 'Palindrome.txt'
csv_file = 'Reverse.csv'
doc_file='Palindrome.docx'

if not os.path.exists(doc_file):
    document=Document()
    document.add_heading('PALINDROME CHECK',0)
    document.save(doc_file)


# create the text file
if not os.path.exists(txt_file):
    with open(txt_file, 'w', newline='') as fp:
        fp.write('PALINDROME LIST' + '\n' + '=======================' + '\n')

# create the csv file
if not os.path.exists(csv_file):
    with open(csv_file, 'w', newline='') as fp:
        writer = csv.writer(fp)
        writer.writerow(['Given', 'Reverse'])


def verify_user_entry(event):
    reverse_label.config(text="")
    if not event.char.isalnum():
        g_input_val.delete(0, tk.END)
        g_input_val.insert(0, g_input_val.get()[::-1])


def clear_all():
    g_input_val.delete(0, tk.END)
    reverse_label.config(text="")


def show_reverse():
    g_string = g_input_val.get()
    # result_string = ""
    if (len(g_string.strip()) > 0):
        result_string = g_string[::-1]
        reverse_label.config(text=result_string)
        record_data(privilege.get())
    else:
        messagebox.showwarning("Empty", "Input is mandatory")
    if (g_string == result_string):
        messagebox.showinfo(result_string, "PALINDROME")
    else:
        messagebox.showinfo(result_string, "NOT A PALINDROME  ")


def record_data(user_type):
    if (user_type == 1):
        g_string = g_input_val.get()
        reversed_str = g_string[::-1]
        if (g_string == reversed_str):
            with open(txt_file, 'a', newline='') as fp:
                fp.write(g_string + '\n')
            doc=Document(doc_file)
            doc.add_paragraph(g_string).bold=True
            doc.save(doc_file)
        else:
            with open(csv_file, 'a', newline='') as fp:
                writer = csv.writer(fp)
                writer.writerow([g_string, reversed_str])

def download_doc():
    result_doc=filedialog.asksaveasfile(
        mode="w",
        defaultextension=".docx",
        filetypes=(("DOCS",".docx"),("All Files",""))
    )
    if result_doc is None:
        return
    name=result_doc.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copyfile(doc_file,target)

def download_csv():
    result_csv=filedialog.asksaveasfile(
        mode="w",
        defaultextension=".csv",
        filetypes=(("CSV",".csv"),("All Files",""))
    )
    if result_csv is None:
        return
    name=result_csv.name
    basename=os.path.basename(name)
    path=os.path.dirname(name)
    target=os.path.join(path,basename)
    shutil.copyfile(csv_file,target)

privilege = tk.IntVar()
window_label = ttk.Label(window, text="Palindrome Check")
window_label.grid(row=0, column=0, sticky=tk.W, padx=20, pady=20)

admin_rdo_button = ttk.Radiobutton(window, text="Administrator", variable=privilege, value=1, command=clear_all)
admin_rdo_button.grid(row=1, column=0, sticky=tk.W, padx=20, pady=20)

user_rdo_button = ttk.Radiobutton(window, text="User", variable=privilege, value=0, command=clear_all)
user_rdo_button.grid(row=1, column=1, sticky=tk.W, padx=20, pady=20)

g_input = ttk.Label(window, text="Enter the string")
g_input.grid(row=2, column=0, sticky=tk.W, padx=20, pady=20)

g_input_val = ttk.Entry(window)
g_input_val.grid(row=2, column=1, sticky=tk.W, padx=20, pady=20)
g_input_val.bind('<KeyRelease>', verify_user_entry)

r_button = ttk.Button(window, text='Reverse', command=show_reverse)
r_button.grid(row=2, column=3, sticky=tk.W, padx=20, pady=20)

g_output = ttk.Label(window, text="Reversed string: ")
g_output.grid(row=3, column=0, sticky=tk.W, padx=20, pady=20)

reverse_label = ttk.Label(window, text="")
reverse_label.grid(row=3, column=1, sticky=tk.W, padx=20, pady=20)

download_label=ttk.Label(window,text="DOWNLOAD")
download_label.grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)

download_doc_btn = ttk.Button(window, text='DOC',command=download_doc)
download_doc_btn.grid(row=4, column=2, sticky=tk.W, padx=5, pady=5)

download_csv_btn = ttk.Button(window, text='CSV',command=download_csv)
download_csv_btn.grid(row=4, column=3, sticky=tk.W, padx=5, pady=5)

window.mainloop()